import re

from DrissionPage import ChromiumPage, SessionPage
from DrissionPage.common import Settings
from DrissionPage.errors import ElementNotFoundError

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Inches
from docx.shared import Pt
from urllib.parse import urlparse, parse_qs

kemu = ['1160566','1161128','1160280','1160279']

page = ChromiumPage()
page.get('https://s.zaixiankaoshi.com/student/114776')
print(page.title)
ele = page.ele('@placeholder=请输入您的学员账号')
ele.input("15615198376")
# 定位到密码文本框并输入密码
page.ele('@placeholder=请输入您的学员密码').input("123456")
# 点击登录按钮
page.ele('tag:button').click()

page.wait.load_start()


# document = Document()
# document.styles['Normal'].font.name = u'宋体'
# document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
# document.styles['Normal'].font.size = Pt(11)

for i in range(5, 6):

    if i == 3:
        continue

    document = Document()
    document.styles['Normal'].font.name = u'宋体'
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    document.styles['Normal'].font.size = Pt(11)

    url = f'https://s.zaixiankaoshi.com/sctk/'
    page.get(url)
    page.wait.eles_loaded('xpath://*[@id="__layout"]/section/section/main/div/div/div[2]/div/div/div/div[1]/div[3]/table/tbody/tr[1]/td[6]/div/button/span')

    a_title =page.ele('xpath://*[@id="__layout"]/section/section/main/div/div/div[2]/div/div/div/div[1]/div[3]/table/tbody/tr['+str(i)+']/td[2]/div/span').text

    buttons = page.ele('xpath://*[@id="__layout"]/section/section/main/div/div/div[2]/div/div/div/div[1]/div[3]/table/tbody/tr['+str(i)+']/td[6]/div/button')
    buttons.click()

    page.wait.eles_loaded('xpath://*[@id="body"]/div/div/div[1]/div[2]/div[3]/p/text()')
    num_t = page.ele('xpath://*[@id="body"]/div/div/div[1]/div[2]/div[3]/div/a[1]/div/p/span').text
    num = int(re.findall(r'\d+', num_t)[0])

    c_url = page.url
    document.add_heading('单选题', level=1)
    page.ele('xpath://*[@id="body"]/div/div/div[1]/div[2]/div[3]/div/a[1]').click()
    page.wait.eles_loaded('xpath://*[@id="body"]/div/div[1]/div[2]/div[1]/div[2]/div[1]/div/div[1]/b/text()')

    # 使用 urlparse() 分割 URL
    parsed_url = urlparse(page.url)
    # 使用 parse_qs() 提取查询参数
    query_params = parse_qs(parsed_url.query)
    seq_num = query_params['sequence'][0]
    while str(num-1) != seq_num:
        title = page.ele('@class=qusetion-box').text
        document.add_paragraph(title)
        option2 = document.add_paragraph('')
        type = page.ele('@class=topic-type').text
        options = page.s_eles('@class^option')
        answer=''
        for opt in options:
            option2.add_run("\n" +opt.raw_text).bold = True
            if 'right' in opt.attrs['class']:
                answer = opt.raw_text.split(" ")[0]
        option2.add_run("\n" + "正确答案: "+answer).bold = True
        page.ele('@@class:el-button el-button--primary el-button--small@@text():下一题', timeout=5).click()
        page.wait(float(1))
        parsed_url = urlparse(page.url)
        # 使用 parse_qs() 提取查询参数
        query_params = parse_qs(parsed_url.query)
        seq_num = query_params['sequence'][0]

    document.add_heading('多选题', level=1)
    page.get(c_url)
    page.wait.eles_loaded('xpath://*[@id="__layout"]/section/section/main/div/div/div[2]/div/div/div/div[1]/div[3]/table/tbody/tr[1]/td[6]/div/button/span')

    num_t2 = page.ele('xpath://*[@id="body"]/div/div/div[1]/div[2]/div[3]/div/a[2]/div/p/span').text
    num2 = int(re.findall(r'\d+', num_t2)[0])

    page.ele('xpath://*[@id="body"]/div/div/div[1]/div[2]/div[3]/div/a[2]').click()
    page.wait.eles_loaded('xpath://*[@id="body"]/div/div[1]/div[2]/div[1]/div[2]/div[1]/div/div[1]/b/text()')
    parsed_url = urlparse(page.url)
    # 使用 parse_qs() 提取查询参数
    query_params = parse_qs(parsed_url.query)
    seq_num = query_params['sequence'][0]
    while str(num2 - 1) != seq_num:
        title = page.ele('@class=qusetion-box').text
        document.add_paragraph(title)
        option2 = document.add_paragraph('')
        type = page.ele('@class=topic-type').text
        options = page.s_eles('@class^option')
        answer = ''
        for opt in options:
            option2.add_run("\n" + opt.raw_text).bold = True
            if 'right' in opt.attrs['class']:
                answer = answer + opt.raw_text.split(" ")[0] + ','
        option2.add_run("\n" + "正确答案: " + answer[:-1]).bold = True
        page.ele('@@class:el-button el-button--primary el-button--small@@text():下一题', timeout=5).click()
        page.wait(float(1))
        parsed_url = urlparse(page.url)
        # 使用 parse_qs() 提取查询参数
        query_params = parse_qs(parsed_url.query)
        seq_num = query_params['sequence'][0]

    document.add_heading('判断题', level=1)
    page.get(c_url)
    page.wait.eles_loaded(
        'xpath://*[@id="__layout"]/section/section/main/div/div/div[2]/div/div/div/div[1]/div[3]/table/tbody/tr[1]/td[6]/div/button/span')

    num_t3 = page.ele('xpath://*[@id="body"]/div/div/div[1]/div[2]/div[3]/div/a[3]/div/p/span').text
    num3 = int(re.findall(r'\d+', num_t3)[0])

    page.ele('xpath://*[@id="body"]/div/div/div[1]/div[2]/div[3]/div/a[3]').click()
    page.wait.eles_loaded('xpath://*[@id="body"]/div/div[1]/div[2]/div[1]/div[2]/div[1]/div/div[1]/b/text()')

    parsed_url = urlparse(page.url)
    # 使用 parse_qs() 提取查询参数
    query_params = parse_qs(parsed_url.query)
    seq_num = query_params['sequence'][0]
    while str(num3 - 1) != seq_num:
        title = page.ele('@class=qusetion-box').text
        document.add_paragraph(title)
        option2 = document.add_paragraph('')
        type = page.ele('@class=topic-type').text
        options = page.s_eles('@class^option')
        answer = ''
        for opt in options:
            option2.add_run("\n" + opt.raw_text).bold = True
            if 'right' in opt.attrs['class']:
                answer = answer + opt.raw_text.split(" ")[0] + ','
        option2.add_run("\n" + "正确答案: " + answer[:-1]).bold = True
        page.ele('@@class:el-button el-button--primary el-button--small@@text():下一题', timeout=5).click()
        page.wait(float(1))
        parsed_url = urlparse(page.url)
        # 使用 parse_qs() 提取查询参数
        query_params = parse_qs(parsed_url.query)
        seq_num = query_params['sequence'][0]

    document.save(a_title+'2.docx')





# for id in kemu:
#     url = f'https://s.zaixiankaoshi.com/select/zxlx/?paperId={id}'
#     page.get(url)
#     type = page.ele('xpath://*[@id="body"]/div/div/div[1]/div[2]/div[3]/div/a[1]/div/p/i').text
#     num = page.ele('xpath://*[@id="body"]/div/div/div[1]/div[2]/div[3]/div/a[1]/div/p/span').text
#     print(type+num)



# page.wait.eles_loaded('xpath://*[@id="body"]/div/div[1]/div[2]/div[1]/div[2]/div[1]/div/div[1]/b/text()')
# title = page.ele('@class=qusetion-box').text
#
#
#
# document.add_heading('单选题', level=1)
# document.add_paragraph(title)
#
# option2 = document.add_paragraph('')
# type = page.ele('@class=topic-type').text
# options = page.s_eles('@class^option')
# for opt in options:
#     option2.add_run(opt.raw_text).bold = True
#
#
# page.ele('@@class:el-button el-button--primary el-button--small@@text():下一题', timeout=5).click()
# document.save('exam_document.docx')



